"""Word 範本編輯工具：read / rename / insert（python-docx based）。

設計：
- 全部 in-place 修改 .docx；呼叫者自行決定是否備份。
- python-docx 在 Jinja 變數 {{ X }} 通常為單一 run，run-level 取代可保留格式；
  若 pattern 跨 run，fallback 改寫整個段落（會 collapse 該段落格式）。
- 所有路徑與副檔名驗證；錯誤回 dict 不拋例外。
"""

import os
import re


def _open_doc(path):
    from docx import Document
    return Document(path)


def _iter_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def _replace_in_paragraph(para, pattern, replacement):
    """先在每個 run 內替換（保留格式）；pattern 跨 run 時 fallback 段落級重寫。"""
    changed = False
    for run in para.runs:
        new = pattern.sub(replacement, run.text)
        if new != run.text:
            run.text = new
            changed = True
    if changed:
        return True

    full_text = "".join(r.text for r in para.runs)
    if not pattern.search(full_text):
        return False

    new_text = pattern.sub(replacement, full_text)
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)
    return True


# ---------- Public APIs ----------


def read_docx_text(path: str, max_paragraphs: int = 0) -> dict:
    """回傳 {paragraphs: list, count: int}（或 error）。"""
    if not path:
        return {"error": "未提供路徑"}
    if not os.path.isfile(path):
        return {"error": f"檔案不存在: {path}"}
    if not path.lower().endswith(".docx"):
        return {"error": "必須是 .docx"}

    try:
        doc = _open_doc(path)
    except Exception as e:
        return {"error": f"無法開啟: {e}"}

    paragraphs = []
    for p in _iter_paragraphs(doc):
        text = p.text
        if text:
            paragraphs.append(text)

    total = len(paragraphs)
    if max_paragraphs and max_paragraphs > 0 and total > max_paragraphs:
        paragraphs = paragraphs[:max_paragraphs]
        truncated = total - max_paragraphs
        paragraphs.append(f"...（其餘 {truncated} 段省略）")

    return {"paragraphs": paragraphs, "count": total}


def rename_template_variable(path: str, old: str, new: str) -> dict:
    """{{ old }} → {{ new }}（容忍空白）。回傳 {changed: int} 或 error。"""
    if not path:
        return {"error": "未提供路徑"}
    if not os.path.isfile(path):
        return {"error": f"檔案不存在: {path}"}
    if not path.lower().endswith(".docx"):
        return {"error": "必須是 .docx"}
    if not old or not new:
        return {"error": "old / new 不可為空"}
    if old == new:
        return {"error": "old 與 new 相同，不需重新命名"}

    pattern = re.compile(r"\{\{\s*" + re.escape(old) + r"\s*\}\}")
    replacement = "{{ " + new + " }}"

    try:
        doc = _open_doc(path)
    except Exception as e:
        return {"error": f"無法開啟: {e}"}

    changed = 0
    for para in _iter_paragraphs(doc):
        if _replace_in_paragraph(para, pattern, replacement):
            changed += 1

    if changed == 0:
        return {"changed": 0, "warning": f"在範本中找不到 {{{{ {old} }}}}（已存檔）"}

    try:
        doc.save(path)
    except Exception as e:
        return {"error": f"存檔失敗: {e}"}

    return {"changed": changed, "from": old, "to": new}


def insert_image_at_anchor(
    path: str,
    anchor: str,
    image_path: str,
    width_mm: int = 80,
) -> dict:
    """在 anchor 文字所在段落的下一行新增段落，並插入圖片。

    用於：圖片檔名 → Word 中對應段落（如「圖 1：流程圖」）後面放圖。
    回傳 {inserted: bool, anchor, image} 或 error。
    """
    if not path:
        return {"error": "未提供 Word 路徑"}
    if not os.path.isfile(path):
        return {"error": f"檔案不存在: {path}"}
    if not path.lower().endswith(".docx"):
        return {"error": "必須是 .docx"}
    if not anchor:
        return {"error": "anchor 不可為空"}
    if not image_path:
        return {"error": "未提供圖片路徑"}
    if not os.path.isfile(image_path):
        return {"error": f"圖片不存在: {image_path}"}

    try:
        w = max(1, int(width_mm))
    except (TypeError, ValueError):
        w = 80

    try:
        doc = _open_doc(path)
    except Exception as e:
        return {"error": f"無法開啟: {e}"}

    target = None
    for para in _iter_paragraphs(doc):
        if anchor in para.text:
            target = para
            break

    if target is None:
        return {"error": f"找不到 anchor: {anchor!r}"}

    try:
        from docx.oxml import OxmlElement
        from docx.shared import Mm
        from docx.text.paragraph import Paragraph

        new_p_elem = OxmlElement("w:p")
        target._p.addnext(new_p_elem)
        new_para = Paragraph(new_p_elem, target._parent)
        run = new_para.add_run()
        run.add_picture(image_path, width=Mm(w))
    except Exception as e:
        return {"error": f"插入圖片失敗: {e}"}

    try:
        doc.save(path)
    except Exception as e:
        return {"error": f"存檔失敗: {e}"}

    return {
        "inserted": True,
        "anchor": anchor,
        "image": image_path,
        "width_mm": w,
    }


def replace_tag_with_image(
    path: str,
    variable: str,
    image_path: str,
    width_mm: int = 80,
) -> dict:
    """把範本中的 {{ variable }} 佔位符,就地換成圖片(同段落原位置)。

    用於「靜態範本欄位」:資料夾圖片依檔名對到範本的 {{圖片欄}},產出前一次插好。
    回傳 {replaced: int, variable, image} 或 error。
    """
    if not path or not os.path.isfile(path):
        return {"error": f"檔案不存在: {path}"}
    if not path.lower().endswith(".docx"):
        return {"error": "必須是 .docx"}
    if not variable:
        return {"error": "variable 不可為空"}
    if not image_path or not os.path.isfile(image_path):
        return {"error": f"圖片不存在: {image_path}"}

    try:
        w = max(1, int(width_mm))
    except (TypeError, ValueError):
        w = 80

    pattern = re.compile(r"\{\{\s*" + re.escape(variable) + r"\s*\}\}")

    try:
        doc = _open_doc(path)
        from docx.shared import Mm
    except Exception as e:
        return {"error": f"無法開啟: {e}"}

    replaced = 0
    for para in _iter_paragraphs(doc):
        full_text = "".join(r.text for r in para.runs)
        if not pattern.search(full_text):
            continue
        # 移除佔位符文字,保留其餘文字於第一個 run,再於段落末尾加圖片
        new_text = pattern.sub("", full_text)
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = new_text
        else:
            para.add_run(new_text)
        try:
            para.add_run().add_picture(image_path, width=Mm(w))
        except Exception as e:
            return {"error": f"插入圖片失敗: {e}"}
        replaced += 1

    if replaced == 0:
        return {"replaced": 0, "warning": f"找不到 {{{{ {variable} }}}}"}

    try:
        doc.save(path)
    except Exception as e:
        return {"error": f"存檔失敗: {e}"}
    return {"replaced": replaced, "variable": variable, "image": image_path, "width_mm": w}


def insert_template_variable(
    path: str,
    anchor: str,
    variable: str,
    position: str = "after",
) -> dict:
    """在 anchor 文字附近插入 {{ variable }}。

    position: "after" / "before" / "replace"
    只處理第一個匹配（避免重複插入）。
    """
    if not path:
        return {"error": "未提供路徑"}
    if not os.path.isfile(path):
        return {"error": f"檔案不存在: {path}"}
    if not path.lower().endswith(".docx"):
        return {"error": "必須是 .docx"}
    if not anchor or not variable:
        return {"error": "anchor / variable 不可為空"}
    if position not in ("after", "before", "replace"):
        return {"error": "position 必須是 after / before / replace"}

    insertion = "{{ " + variable + " }}"

    try:
        doc = _open_doc(path)
    except Exception as e:
        return {"error": f"無法開啟: {e}"}

    inserted = False
    for para in _iter_paragraphs(doc):
        full_text = "".join(r.text for r in para.runs)
        if anchor not in full_text:
            continue

        if position == "after":
            new_text = full_text.replace(anchor, anchor + insertion, 1)
        elif position == "before":
            new_text = full_text.replace(anchor, insertion + anchor, 1)
        else:  # replace
            new_text = full_text.replace(anchor, insertion, 1)

        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = new_text
        else:
            para.add_run(new_text)
        inserted = True
        break

    if not inserted:
        return {"error": f"找不到 anchor: {anchor!r}"}

    try:
        doc.save(path)
    except Exception as e:
        return {"error": f"存檔失敗: {e}"}

    return {
        "inserted": True,
        "anchor": anchor,
        "variable": variable,
        "position": position,
    }
